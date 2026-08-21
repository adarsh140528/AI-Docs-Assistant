import os
from pypdf import PdfReader
import docx


def clean_text(text: str) -> str:
    """Clean invalid Unicode surrogate characters and normalize whitespace."""
    if not text:
        return ""
    # Remove invalid Unicode surrogate characters
    text = text.encode("utf-16", "surrogatepass").decode("utf-16", "ignore")
    # Clean UTF-8
    text = text.encode("utf-8", errors="ignore").decode("utf-8")
    return text.strip()


class DocLoader:
    """
    Multi-format document loader supporting:
    - .pdf (via pypdf)
    - .docx (via python-docx)
    - .txt / .md (plain text)
    """

    def __init__(self, file_path: str, original_filename: str = None):
        self.file_path = file_path
        self.original_filename = original_filename or os.path.basename(file_path)

    def load(self) -> list[dict]:
        """
        Extract text from file and return list of page-level dictionaries:
        [{"page": int, "text": str, "source": str}, ...]
        """
        ext = os.path.splitext(self.file_path)[1].lower()

        if ext == ".pdf":
            return self._load_pdf()
        elif ext in [".docx", ".doc"]:
            return self._load_docx()
        elif ext in [".txt", ".md", ".csv", ".json"]:
            return self._load_text()
        else:
            # Fallback to plain text
            return self._load_text()

    def _load_pdf(self) -> list[dict]:
        pages = []
        try:
            reader = PdfReader(self.file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    cleaned = clean_text(text)
                    if cleaned:
                        pages.append({
                            "page": i + 1,
                            "text": cleaned,
                            "source": self.original_filename
                        })
        except Exception as e:
            print(f"Error loading PDF {self.file_path}: {e}")
        return pages

    def _load_docx(self) -> list[dict]:
        pages = []
        try:
            doc = docx.Document(self.file_path)
            full_text = []
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)

            combined_text = clean_text("\n\n".join(full_text))
            if combined_text:
                pages.append({
                    "page": 1,
                    "text": combined_text,
                    "source": self.original_filename
                })
        except Exception as e:
            print(f"Error loading DOCX {self.file_path}: {e}")
        return pages

    def _load_text(self) -> list[dict]:
        pages = []
        try:
            with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = clean_text(f.read())
            if content:
                pages.append({
                    "page": 1,
                    "text": content,
                    "source": self.original_filename
                })
        except Exception as e:
            print(f"Error loading text file {self.file_path}: {e}")
        return pages
